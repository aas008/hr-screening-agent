"""
GitHub Resume Loader
Handles loading and parsing resume files from GitHub repository
"""

import os
import re
import io
import requests
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass
from datetime import datetime
import PyPDF2
import docx

@dataclass
class CandidateInfo:
    """Data class for candidate information"""
    name: str
    email: str
    phone: str
    resume_text: str
    file_name: str
    application_date: str
    raw_file_size: int = 0

class GitHubResumeLoader:
    """Load and parse resumes from GitHub repository"""
    
    def __init__(self, github_token: str, repo_owner: str, repo_name: str):
        self.github_token = github_token
        self.repo_owner = repo_owner
        self.repo_name = repo_name
        self.base_url = f"https://api.github.com/repos/{repo_owner}/{repo_name}/contents"
        self.session = requests.Session()
        self.session.headers.update({
            "Authorization": f"token {github_token}",
            "Accept": "application/vnd.github.v3+json"
        })
    
    def load_resumes_from_job_role(self, job_role: str) -> Tuple[List[CandidateInfo], List[str]]:
        """
        Load all resume files from /resumes/active/{job_role}/ folder
        Returns: (successful_candidates, error_messages)
        """
        
        folder_path = f"resumes/active/{job_role}"
        print(f"📂 Loading resumes from GitHub: /{folder_path}/")
        
        try:
            # Get folder contents from GitHub API
            url = f"{self.base_url}/{folder_path}"
            response = self.session.get(url)
            
            if response.status_code == 404:
                error_msg = f"Folder not found: /{folder_path}/"
                print(f"❌ {error_msg}")
                return [], [f"{error_msg}. Please create the folder and add resume files."]
            
            response.raise_for_status()
            files = response.json()
            
            # Filter for resume files
            resume_files = [
                f for f in files 
                if f['type'] == 'file' and f['name'].lower().endswith(('.pdf', '.docx', '.doc'))
            ]
            
            if not resume_files:
                error_msg = f"No resume files found in /{folder_path}/"
                print(f"⚠️  {error_msg}")
                return [], [f"{error_msg}. Please add PDF or DOCX files."]
            
            print(f"✅ Found {len(resume_files)} resume files")
            
            candidates = []
            errors = []
            
            for i, file_info in enumerate(resume_files, 1):
                print(f"   📄 Processing {i}/{len(resume_files)}: {file_info['name']}")
                
                try:
                    candidate = self._process_resume_file(file_info)
                    if candidate:
                        candidates.append(candidate)
                        print(f"   ✅ Extracted: {candidate.name} ({candidate.email})")
                    else:
                        error_msg = f"Could not extract candidate info from {file_info['name']}"
                        print(f"   ⚠️  {error_msg}")
                        errors.append(error_msg)
                        
                except Exception as e:
                    error_msg = f"Error processing {file_info['name']}: {str(e)}"
                    print(f"   ❌ {error_msg}")
                    errors.append(error_msg)
            
            print(f"📊 Successfully processed {len(candidates)} candidates")
            if errors:
                print(f"⚠️  {len(errors)} files had processing errors")
            
            return candidates, errors
            
        except requests.RequestException as e:
            error_msg = f"GitHub API error: {str(e)}"
            print(f"❌ {error_msg}")
            return [], [error_msg]
    
    def _process_resume_file(self, file_info: Dict) -> Optional[CandidateInfo]:
        """Process a single resume file from GitHub"""
        
        try:
            # Download file content
            file_content = self._download_file_content(file_info['download_url'])
            file_size = len(file_content)
            
            # Parse resume text
            resume_text = self._parse_resume_file(file_content, file_info['name'])
            
            if not resume_text or len(resume_text.strip()) < 50:
                print(f"   ⚠️  Insufficient text extracted from {file_info['name']}")
                return None
            
            # Extract candidate information
            candidate = self._extract_candidate_info(resume_text, file_info['name'], file_size)
            
            return candidate
            
        except Exception as e:
            print(f"   ❌ Processing error: {e}")
            return None
    
    def _download_file_content(self, download_url: str) -> bytes:
        """Download file content from GitHub"""
        
        response = self.session.get(download_url)
        response.raise_for_status()
        return response.content
    
    def _parse_resume_file(self, file_content: bytes, filename: str) -> str:
        """Parse PDF or DOCX file to extract text"""
        
        try:
            if filename.lower().endswith('.pdf'):
                return self._extract_pdf_text(file_content)
            elif filename.lower().endswith(('.docx', '.doc')):
                return self._extract_docx_text(file_content)
            else:
                return ""
        except Exception as e:
            raise Exception(f"Failed to parse {filename}: {e}")
    
    def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF using PyPDF2"""
        
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            if len(pdf_reader.pages) == 0:
                raise Exception("PDF has no pages")
            
            text = ""
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
                except Exception as e:
                    print(f"   ⚠️  Warning: Could not extract text from page {page_num + 1}: {e}")
            
            if not text.strip():
                raise Exception("No text could be extracted from PDF")
            
            return text.strip()
            
        except Exception as e:
            raise Exception(f"PDF parsing failed: {e}")
    
    def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX using python-docx"""
        
        try:
            doc_file = io.BytesIO(file_content)
            doc = docx.Document(doc_file)
            
            text_parts = []
            
            # Extract paragraph text
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extract table text
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            if not text_parts:
                raise Exception("No text found in DOCX document")
            
            return "\n".join(text_parts)
            
        except Exception as e:
            raise Exception(f"DOCX parsing failed: {e}")
    
    def _extract_candidate_info(self, resume_text: str, filename: str, file_size: int) -> Optional[CandidateInfo]:
        """Extract structured candidate information from resume text"""
        
        try:
            # Extract name
            name = self._extract_name(resume_text, filename)
            if not name:
                print(f"   ⚠️  Could not extract name from {filename}")
                return None
            
            # Extract email
            email = self._extract_email(resume_text)
            if not email:
                print(f"   ⚠️  Could not extract email from {filename}")
                return None
            
            # Extract phone (optional)
            phone = self._extract_phone(resume_text) or "Not provided"
            
            return CandidateInfo(
                name=name,
                email=email,
                phone=phone,
                resume_text=resume_text,
                file_name=filename,
                application_date=datetime.now().isoformat(),
                raw_file_size=file_size
            )
            
        except Exception as e:
            print(f"   ❌ Info extraction error: {e}")
            return None
    
    def _extract_name(self, text: str, filename: str) -> Optional[str]:
        """Extract candidate name from resume text"""
        
        # Method 1: Try filename first (most reliable)
        name_from_file = self._clean_filename_for_name(filename)
        if self._is_valid_name(name_from_file):
            return name_from_file
        
        # Method 2: Look for name patterns in first few lines
        lines = [line.strip() for line in text.split('\n')[:10] if line.strip()]
        
        for line in lines:
            # Skip common resume headers
            if any(keyword in line.lower() for keyword in [
                'resume', 'cv', 'curriculum', 'vitae', 'profile', 'summary',
                'contact', 'phone', 'email', '@', 'www', 'http'
            ]):
                continue
            
            # Look for name-like patterns
            cleaned_line = re.sub(r'[^\w\s]', '', line).strip()
            if self._is_valid_name(cleaned_line):
                return cleaned_line
        
        # Method 3: Look for patterns like "Name: John Smith"
        name_patterns = [
            r'name\s*[:\-]\s*([a-z\s]{2,30})',
            r'^([A-Z][a-z]+ [A-Z][a-z]+(?:\s[A-Z][a-z]+)?)\s*$',
            r'candidate\s*[:\-]\s*([a-z\s]{2,30})'
        ]
        
        for pattern in name_patterns:
            matches = re.findall(pattern, text[:500], re.IGNORECASE | re.MULTILINE)
            for match in matches:
                cleaned_match = match.strip()
                if self._is_valid_name(cleaned_match):
                    return cleaned_match
        
        # Fallback to cleaned filename
        return name_from_file if name_from_file else None
    
    def _clean_filename_for_name(self, filename: str) -> str:
        """Clean filename to extract potential name"""
        
        # Remove extension
        name = filename.split('.')[0]
        
        # Remove common resume keywords
        remove_keywords = ['resume', 'cv', 'curriculum', 'vitae', '_resume', '-resume']
        for keyword in remove_keywords:
            name = re.sub(keyword, '', name, flags=re.IGNORECASE)
        
        # Replace separators with spaces
        name = re.sub(r'[_\-\.]+', ' ', name)
        
        # Clean up spacing and capitalize
        name = ' '.join(word.capitalize() for word in name.split() if word.strip())
        
        return name.strip()
    
    def _is_valid_name(self, name: str) -> bool:
        """Check if extracted text looks like a valid name"""
        
        if not name or len(name) < 3:
            return False
        
        words = name.split()
        
        # Should have 2-4 words
        if not (2 <= len(words) <= 4):
            return False
        
        # Each word should be reasonable length
        if any(len(word) < 2 or len(word) > 20 for word in words):
            return False
        
        # Should contain only letters and spaces
        if not re.match(r'^[A-Za-z\s]+$', name):
            return False
        
        # Each word should start with capital letter
        if not all(word[0].isupper() for word in words):
            return False
        
        return True
    
    def _extract_email(self, text: str) -> Optional[str]:
        """Extract email address from resume text"""
        
        # Common email patterns
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        
        if not emails:
            return None
        
        # Filter out common non-personal emails
        filtered_emails = []
        for email in emails:
            email_lower = email.lower()
            # Skip obviously non-personal emails
            if not any(keyword in email_lower for keyword in [
                'noreply', 'admin', 'info', 'contact', 'support', 'sales',
                'hr', 'jobs', 'recruiting', 'company', 'example'
            ]):
                filtered_emails.append(email)
        
        # Return first valid email
        return filtered_emails[0] if filtered_emails else emails[0]
    
    def _extract_phone(self, text: str) -> Optional[str]:
        """Extract phone number from resume text"""
        
        # Phone number patterns
        phone_patterns = [
            r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',  # US format
            r'\+\d{1,3}[-.\s]?\d{3,4}[-.\s]?\d{3,4}[-.\s]?\d{3,4}',  # International
            r'\d{3}[-.\s]?\d{3}[-.\s]?\d{4}',  # Simple format
        ]
        
        for pattern in phone_patterns:
            phones = re.findall(pattern, text)
            if phones:
                # Clean up the phone number
                phone = phones[0]
                # Remove common formatting
                phone = re.sub(r'[^\d+]', '', phone)
                # Format nicely
                if phone.startswith('+'):
                    return phone
                elif len(phone) == 10:
                    return f"({phone[:3]}) {phone[3:6]}-{phone[6:]}"
                else:
                    return phone
        
        return None

def validate_github_config(config: Dict) -> bool:
    """Validate GitHub configuration by testing API access"""
    
    try:
        # Test GitHub API access
        loader = GitHubResumeLoader(
            config['github_token'],
            config['repo_owner'],
            config['repo_name']
        )
        
        # Test repository access
        url = f"{loader.base_url}/README.md"
        response = loader.session.get(url)
        
        if response.status_code == 401:
            print("❌ GitHub token is invalid or expired")
            return False
        elif response.status_code == 404:
            print("❌ Repository not found or not accessible")
            print(f"   Repository: {config['repo_owner']}/{config['repo_name']}")
            return False
        elif response.status_code != 200:
            print(f"❌ GitHub API error: {response.status_code}")
            return False
        else:
            print("✅ GitHub configuration validated successfully")
            return True
            
    except Exception as e:
        print(f"❌ GitHub validation error: {e}")
        return False

def test_resume_loading():
    """Test function for resume loading (development use)"""
    
    import os
    from dotenv import load_dotenv
    
    load_dotenv()
    
    config = {
        'github_token': os.getenv('GITHUB_TOKEN'),
        'repo_owner': os.getenv('GITHUB_REPO_OWNER'),
        'repo_name': os.getenv('GITHUB_REPO_NAME')
    }
    
    if not all(config.values()):
        print("❌ Missing environment variables for testing")
        return
    
    loader = GitHubResumeLoader(**config)
    candidates, errors = loader.load_resumes_from_job_role('react-developer')
    
    print(f"\n📊 TEST RESULTS:")
    print(f"   Candidates loaded: {len(candidates)}")
    print(f"   Errors encountered: {len(errors)}")
    
    for candidate in candidates:
        print(f"   • {candidate.name} ({candidate.email}) - {len(candidate.resume_text)} chars")
    
    for error in errors:
        print(f"   ⚠️  {error}")

if __name__ == "__main__":
    test_resume_loading()